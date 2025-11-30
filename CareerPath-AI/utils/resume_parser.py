import io
import re
from typing import Dict, List, Optional, Tuple
import PyPDF2
import pdfplumber
from docx import Document


class ResumeParser:
    """
    Resume parser that extracts text and structured information from PDF and DOCX files.
    """
    
    SKILL_KEYWORDS = [
        # Programming Languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust', 'php', 'swift', 'kotlin', 'scala', 'r',
        # Web Technologies
        'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi', 'spring', 'asp.net',
        # Data Science & ML
        'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'matplotlib',
        'data analysis', 'data visualization', 'nlp', 'computer vision', 'neural networks', 'ai', 'artificial intelligence',
        # Databases
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'oracle', 'sqlite',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'ci/cd', 'terraform', 'ansible',
        # Tools & Frameworks
        'git', 'github', 'gitlab', 'jira', 'agile', 'scrum', 'rest api', 'graphql', 'microservices',
        # Soft Skills
        'leadership', 'communication', 'problem solving', 'teamwork', 'project management', 'analytical',
        # Certifications
        'aws certified', 'azure certified', 'pmp', 'scrum master', 'data science', 'machine learning engineer'
    ]
    
    SECTION_PATTERNS = {
        'experience': r'(?i)(work\s*experience|professional\s*experience|employment\s*history|experience)',
        'education': r'(?i)(education|academic\s*background|qualifications)',
        'skills': r'(?i)(skills|technical\s*skills|competencies|expertise)',
        'projects': r'(?i)(projects|personal\s*projects|portfolio|key\s*projects)',
        'certifications': r'(?i)(certifications|certificates|accreditations)',
        'summary': r'(?i)(summary|profile|objective|about\s*me)'
    }

    def __init__(self):
        pass

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using multiple methods for reliability."""
        text = ""
        
        # Try pdfplumber first (more reliable for structured text)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
        
        # Fallback to PyPDF2 if pdfplumber fails or returns empty
        if not text.strip():
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception:
                pass
        
        return text.strip()

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from resume file based on file type."""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            text = self.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            text = self.extract_text_from_docx(file_content)
        elif filename_lower.endswith('.doc'):
            raise ValueError("DOC format not supported. Please use DOCX or PDF.")
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
        
        if not text.strip():
            raise ValueError("Could not extract text from the resume. The file may be corrupted or empty.")
        
        return text

    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.SKILL_KEYWORDS:
            # Use word boundaries for accurate matching
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill.title())
        
        return list(set(found_skills))

    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from the resume."""
        sections = {}
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check if line is a section header
            found_section = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped) and len(line_stripped) < 50:
                    found_section = section_name
                    break
            
            if found_section:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = found_section
                current_content = []
            else:
                current_content.append(line_stripped)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections

    def extract_projects(self, text: str) -> List[Dict[str, str]]:
        """Extract project information from resume."""
        projects = []
        sections = self.extract_sections(text)
        
        project_text = sections.get('projects', '')
        if not project_text:
            # Try to find projects in the full text
            project_match = re.search(
                r'(?i)projects?\s*[:\n](.+?)(?=(?:experience|education|skills|certifications|$))',
                text,
                re.DOTALL
            )
            if project_match:
                project_text = project_match.group(1)
        
        if project_text:
            # Split by bullet points or new lines with capital letters
            project_entries = re.split(r'\n(?=[A-Z•\-\*])', project_text)
            
            for entry in project_entries:
                entry = entry.strip()
                if len(entry) > 20:  # Minimum length for a valid project entry
                    # Try to extract project name and description
                    lines = entry.split('\n')
                    name = lines[0].strip('•-* ')
                    description = ' '.join(lines[1:]).strip() if len(lines) > 1 else ''
                    
                    if name:
                        projects.append({
                            'name': name[:100],  # Limit name length
                            'description': description[:500] if description else 'No description available'
                        })
        
        return projects[:10]  # Limit to 10 projects

    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume."""
        contact = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None
        }
        
        # Email pattern
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone pattern
        phone_match = re.search(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}', text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group()
        
        # GitHub pattern
        github_match = re.search(r'github\.com/[\w-]+', text, re.IGNORECASE)
        if github_match:
            contact['github'] = github_match.group()
        
        return contact

    def parse_resume(self, file_content: bytes, filename: str) -> Dict:
        """
        Main method to parse a resume and extract all information.
        Returns a dictionary with all extracted data.
        """
        # Extract raw text
        text = self.extract_text(file_content, filename)
        
        # Extract structured information
        skills = self.extract_skills(text)
        sections = self.extract_sections(text)
        projects = self.extract_projects(text)
        contact = self.extract_contact_info(text)
        
        return {
            'raw_text': text,
            'skills': skills,
            'sections': sections,
            'projects': projects,
            'contact': contact,
            'word_count': len(text.split()),
            'has_projects': len(projects) > 0
        }
